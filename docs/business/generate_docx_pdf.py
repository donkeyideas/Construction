"""
Generate Buildwrk White Paper as Word (.docx) and PDF
Matches the homepage design system colors and typography.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ─── Brand Colors ───
BLUE = RGBColor(0x1D, 0x4E, 0xD8)
AMBER = RGBColor(0xB4, 0x53, 0x09)
TEXT_DARK = RGBColor(0x29, 0x25, 0x24)
MUTED = RGBColor(0x78, 0x71, 0x6C)
WHITE_COLOR = RGBColor(0xFF, 0xFF, 0xFF)
BG_WARM = RGBColor(0xF5, 0xF0, 0xEB)

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_row(table, cells_data, header=False):
    """Add a row to a table with styled cells."""
    row = table.add_row()
    for i, (text, bold, color) in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9 if not header else 9)
        run.font.bold = bold
        run.font.color.rgb = color
        run.font.name = "Calibri"
        if header:
            set_cell_shading(cell, "1D4ED8")
            run.font.color.rgb = WHITE_COLOR
    return row

# ═══════════════════════════════════════════════════════════
# WORD DOCUMENT
# ═══════════════════════════════════════════════════════════

doc = Document()

# ── Set default font ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = TEXT_DARK

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ═══════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(120)
run = p.add_run("Buildwrk")
run.font.size = Pt(48)
run.font.color.rgb = BLUE
run.font.bold = True
run.font.name = "Georgia"

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("The Unified Construction Intelligence Platform")
run.font.size = Pt(20)
run.font.color.rgb = TEXT_DARK
run.font.name = "Georgia"

# Amber divider
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("_" * 15)
run.font.color.rgb = AMBER
run.font.size = Pt(14)

# White paper subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(20)
run = p.add_run("A White Paper on Modernizing the\n$1.3 Trillion Construction Industry")
run.font.size = Pt(16)
run.font.color.rgb = MUTED
run.font.name = "Georgia"

# Date
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(40)
run = p.add_run("February 2026")
run.font.size = Pt(14)
run.font.color.rgb = AMBER
run.font.bold = True

# Confidential
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(60)
run = p.add_run("CONFIDENTIAL")
run.font.size = Pt(10)
run.font.color.rgb = MUTED
run.font.all_caps = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════

p = doc.add_paragraph()
run = p.add_run("Table of Contents")
run.font.size = Pt(24)
run.font.color.rgb = BLUE
run.font.bold = True
run.font.name = "Georgia"

doc.add_paragraph("")

toc_items = [
    ("1.", "Executive Summary", 3),
    ("2.", "The Problem: Fragmented Technology in Construction", 3),
    ("", "2.1  The Software Stack Tax", 4),
    ("", "2.2  The Productivity Gap", 4),
    ("", "2.3  The Mid-Market Vacuum", 5),
    ("3.", "The Solution: Unified Data Architecture", 5),
    ("", "3.1  Single Source of Truth", 5),
    ("", "3.2  Multi-Tenant Architecture", 6),
    ("", "3.3  GAAP-Compliant Financial Engine", 6),
    ("", "3.4  AI-Powered Intelligence Layer", 7),
    ("4.", "Technical Architecture", 8),
    ("5.", "Market Opportunity", 9),
    ("6.", "Business Model", 10),
    ("7.", "Conclusion", 11),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    indent = "     " if num == "" else ""
    run = p.add_run(f"{indent}{num} {title}")
    run.font.size = Pt(11 if num else 10)
    run.font.color.rgb = TEXT_DARK if num else MUTED
    run.font.bold = bool(num)
    # Right-aligned page number
    run2 = p.add_run(f"\t{page}")
    run2.font.color.rgb = MUTED
    run2.font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════

def add_heading_styled(text, level=1):
    """Add a branded heading."""
    p = doc.add_paragraph()
    p.space_before = Pt(24 if level == 1 else 16)
    p.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(24 if level == 1 else 18 if level == 2 else 14)
    run.font.color.rgb = BLUE if level <= 2 else AMBER
    run.font.bold = True
    run.font.name = "Georgia"
    return p

def add_body(text, bold=False, color=None):
    """Add body text."""
    p = doc.add_paragraph()
    p.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = color or TEXT_DARK
    run.font.bold = bold
    run.font.name = "Calibri"
    return p

def add_bullet(text, bold_prefix=""):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK
    return p

# ═══════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════

add_heading_styled("1. Executive Summary")

add_body("The construction industry generates $1.3 trillion in annual revenue in the United States alone, yet remains one of the least digitized sectors of the global economy. General contractors, developers, and property managers still rely on fragmented technology stacks \u2014 using one tool for project management, another for accounting, a third for safety compliance, and spreadsheets to bridge the gaps between them.")

add_body("Buildwrk is a unified construction ERP platform that consolidates project management, GAAP-compliant financial accounting, property management, safety compliance, equipment tracking, CRM, and AI-powered analytics into a single, multi-tenant SaaS application. Built on modern cloud architecture (Next.js, Supabase/PostgreSQL, Vercel), Buildwrk eliminates the data silos that cost construction companies an estimated 5-10% of project revenue in inefficiency.")

add_body("This white paper examines the structural problems in construction technology, the architectural decisions behind Buildwrk, and the economic case for a unified platform approach.")

# ═══════════════════════════════════════════════════════════
# THE PROBLEM
# ═══════════════════════════════════════════════════════════

add_heading_styled("2. The Problem: Fragmented Technology in Construction")

add_heading_styled("2.1 The Software Stack Tax", level=2)

add_body("A typical mid-market general contractor ($20M-$200M revenue) operates with 5-8 separate software systems:")

# Table: Software stack
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# Header row
for i, header in enumerate(["Function", "Common Tool", "Monthly Cost", "Data Silo"]):
    cell = table.rows[0].cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(header)
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = WHITE_COLOR
    run.font.name = "Calibri"
    set_cell_shading(cell, "1D4ED8")

rows_data = [
    ["Project Management", "Procore", "$5,000-$9,000", "Schedules, RFIs, daily logs"],
    ["Accounting", "Sage 300 CRE", "$3,000-$8,000", "GL, AP/AR, job costing"],
    ["Estimating", "PlanSwift/HCSS", "$1,000-$3,000", "Bid data, takeoffs"],
    ["Safety", "SafetyCulture", "$500-$2,000", "Incidents, inspections"],
    ["HR/Payroll", "ADP/Paychex", "$1,000-$3,000", "Timesheets, wages"],
    ["Document Mgmt", "Box/SharePoint", "$500-$1,500", "Plans, specs, submittals"],
    ["CRM", "Salesforce", "$1,500-$5,000", "Leads, bids, contacts"],
]

for row_data in rows_data:
    row = table.add_row()
    for i, text in enumerate(row_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = TEXT_DARK
        run.font.name = "Calibri"

doc.add_paragraph("")
add_body("Total: $12,500 --$31,500/month ($150K-$378K/year)", bold=True)

add_body("Beyond direct cost, these systems don't share data natively. When a change order is approved in Procore, someone must manually create a journal entry in Sage. When a safety incident occurs, there's no automatic link to the project budget. When a payment is received, the project manager doesn't see updated AR aging in real-time.")

add_heading_styled("2.2 The Productivity Gap", level=2)

add_body("McKinsey & Company has repeatedly identified construction as the industry with the lowest productivity growth over the past 50 years. While manufacturing productivity has grown 760% since 1970, construction has remained essentially flat.")

add_body("The root cause is not lack of technology \u2014 it's lack of integrated technology. Data enters the ecosystem once, then gets re-entered, reformatted, and reconciled across systems. A single change order can trigger 6-8 manual data entry steps across project management, accounting, and documentation systems.")

add_heading_styled("2.3 The Mid-Market Vacuum", level=2)

add_body("Enterprise players (Oracle/Primavera, Trimble, Autodesk Construction Cloud) serve firms above $500M in revenue with complex, expensive implementations. Small-business tools (Contractor Foreman, CoConstruct) serve firms below $5M with simplified features.")

add_body("The mid-market ($5M-$500M) \u2014 which represents the majority of construction firms by count and a significant share of revenue \u2014 is underserved. These firms need enterprise-grade features (GAAP accounting, multi-project management, safety compliance) at a price point that doesn't require a six-figure annual commitment.")

# ═══════════════════════════════════════════════════════════
# THE SOLUTION
# ═══════════════════════════════════════════════════════════

add_heading_styled("3. The Solution: Unified Data Architecture")

add_heading_styled("3.1 Single Source of Truth", level=2)

add_body("Buildwrk's fundamental design principle is that every piece of construction data \u2014 from a daily log entry to a journal entry line \u2014 lives in a single relational database with enforced referential integrity. There are no sync engines, middleware layers, or eventual-consistency models between modules.")

add_body("When a change order is approved, the system:", bold=True)

add_bullet("Updates the project budget (Project Management module)")
add_bullet("Creates a balanced journal entry: DR AR / CR Revenue for owner-initiated COs, or DR Expense / CR AP for cost COs (Financial module)")
add_bullet("Adjusts the contract value (Contract Management module)")
add_bullet("Logs the approval in the audit trail (Compliance module)")
add_bullet("Notifies the project manager and accountant (Communication module)")

add_body("All five operations execute against the same database in a single transaction. There is no delay, no manual intervention, and no reconciliation needed.")

add_heading_styled("3.2 Multi-Tenant Architecture with Enterprise Security", level=2)

add_body("Buildwrk uses PostgreSQL Row-Level Security (RLS) to enforce data isolation between companies. Every table includes a company_id column, and 461 RLS policies ensure that:")

add_bullet("Company A can never see Company B's data, even through a SQL injection attempt")
add_bullet("Within a company, 7 role levels (Owner, Admin, Project Manager, Superintendent, Accountant, Field Worker, Viewer) enforce least-privilege access")
add_bullet("Tenant, vendor, and employee portals provide scoped external access without exposing internal data")

add_heading_styled("3.3 GAAP-Compliant Financial Engine", level=2)

add_body("Construction accounting has unique requirements that general-purpose tools like QuickBooks cannot adequately address:")

add_body("Retainage tracking", bold=True)
add_body("Construction contracts typically hold 5-10% of each payment until project completion. Buildwrk creates separate journal entry lines for retainage (DR Retainage Receivable / CR AR), tracking these amounts independently from standard AR/AP aging.")

add_body("Change order accounting", bold=True)
add_body("Owner-initiated change orders increase contract revenue; cost change orders increase project expense. Buildwrk automatically generates the correct journal entries based on change order type upon approval.")

add_body("Job costing by CSI code", bold=True)
add_body("The Construction Specifications Institute (CSI) MasterFormat is the industry standard for categorizing work. Buildwrk's budget module supports CSI-code-level cost tracking with budgeted-vs-actual variance analysis.")

add_body("The financial engine generates all four primary financial statements (Income Statement, Balance Sheet, Cash Flow Statement, Trial Balance) plus AR/AP aging, general ledger detail, and job cost reports \u2014 all from the same underlying journal entry data.")

add_heading_styled("3.4 AI-Powered Intelligence Layer", level=2)

add_body("Buildwrk integrates 9 LLM providers (OpenAI, Anthropic Claude, Google Gemini, Groq, Mistral, Cohere, DeepSeek, xAI Grok, AWS Bedrock) through a unified provider router. Companies can:")

add_bullet("Choose their preferred AI provider based on cost, capability, or data residency requirements")
add_bullet("Set per-provider monthly budget limits")
add_bullet("Use AI for natural-language queries against live project data")
add_bullet("Generate reports with AI-synthesized insights")
add_bullet("Track token usage and costs per user")

add_body("The AI system uses function calling to query real company data \u2014 it's not generic chatbot responses, but actual database queries executed in the context of the user's company, projects, and permissions.")

# ═══════════════════════════════════════════════════════════
# TECHNICAL ARCHITECTURE
# ═══════════════════════════════════════════════════════════

add_heading_styled("4. Technical Architecture")

add_heading_styled("4.1 Technology Stack", level=2)

# Tech stack table
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

for i, header in enumerate(["Layer", "Technology", "Rationale"]):
    cell = table.rows[0].cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(header)
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = WHITE_COLOR
    set_cell_shading(cell, "1D4ED8")

tech_rows = [
    ["Frontend", "Next.js 16.1.6 (App Router)", "Server-first rendering, optimal SEO"],
    ["Language", "TypeScript (strict mode)", "Type safety across 646 source files"],
    ["Database", "PostgreSQL (Supabase)", "ACID compliance, RLS, real-time"],
    ["Auth", "Supabase Auth", "Email/password, OAuth-ready"],
    ["Hosting", "Vercel", "Serverless auto-scaling, global CDN"],
    ["AI", "Vercel AI SDK", "Streaming, function calling, 9 providers"],
    ["Styling", "Custom CSS (43 modules)", "Design tokens, zero-runtime overhead"],
]

for row_data in tech_rows:
    row = table.add_row()
    for i, text in enumerate(row_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = TEXT_DARK

add_heading_styled("4.2 Platform Metrics", level=2)

metrics = [
    ("80,000+", "lines of production code (TypeScript, SQL, CSS)"),
    ("646", "TypeScript source files with 100% type coverage"),
    ("46", "database tables with full relational schema"),
    ("461", "Row-Level Security policies for multi-tenant isolation"),
    ("177", "REST API endpoints across 20+ feature areas"),
    ("178", "application pages covering 19 feature modules"),
    ("35", "progressive database migrations"),
    ("9", "AI/LLM providers with 30+ supported models"),
    ("13", "complete demo datasets covering all major construction verticals"),
]

for number, description in metrics:
    p = doc.add_paragraph()
    run = p.add_run(f"{number} ")
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    run.font.bold = True
    run = p.add_run(description)
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_DARK

# ═══════════════════════════════════════════════════════════
# MARKET OPPORTUNITY
# ═══════════════════════════════════════════════════════════

add_heading_styled("5. Market Opportunity")

add_heading_styled("5.1 Market Size", level=2)

add_body("The global construction management software market is valued at $11.58 billion in 2026, growing at 8.88% CAGR to reach $17.72 billion by 2031 (Mordor Intelligence). The broader construction software market is projected to reach $21.04 billion by 2032 at 10.9% CAGR.")

add_heading_styled("5.2 Competitive Landscape", level=2)

# Competitor table
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

for i, header in enumerate(["Competitor", "Annual Cost (25 users)", "Focus", "Key Weakness"]):
    cell = table.rows[0].cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(header)
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = WHITE_COLOR
    set_cell_shading(cell, "1D4ED8")

comp_rows = [
    ["Procore", "$60K-$112K", "Project Management", "No native accounting"],
    ["Sage 300 CRE", "$36K-$96K", "Accounting", "Legacy UI, no PM"],
    ["Buildertrend", "$6K-$10K", "Residential", "Weak financials"],
    ["Viewpoint", "$50K-$150K", "ERP", "Complex implementation"],
    ["Contractor Foreman", "$3K-$6K", "Small business", "Limited depth"],
    ["Buildwrk", "$3K-$6K", "Unified", "Building brand"],
]

for row_data in comp_rows:
    row = table.add_row()
    for i, text in enumerate(row_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = TEXT_DARK
        if row_data[0] == "Buildwrk":
            run.font.bold = True
            run.font.color.rgb = BLUE

# ═══════════════════════════════════════════════════════════
# BUSINESS MODEL
# ═══════════════════════════════════════════════════════════

add_heading_styled("6. Business Model")

add_heading_styled("6.1 SaaS Subscription Tiers", level=2)

table = doc.add_table(rows=1, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

for i, header in enumerate(["Tier", "Price", "Target", "Users", "Projects"]):
    cell = table.rows[0].cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(header)
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = WHITE_COLOR
    set_cell_shading(cell, "1D4ED8")

pricing_rows = [
    ["Starter", "$99/mo", "Solo GCs, small subs", "5", "3 active"],
    ["Professional", "$249/mo", "Mid-size GCs", "25", "15 active"],
    ["Enterprise", "$499/mo", "Large GCs, developers", "Unlimited", "Unlimited"],
]

for row_data in pricing_rows:
    row = table.add_row()
    for i, text in enumerate(row_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = TEXT_DARK

add_heading_styled("6.2 Go-to-Market Strategy", level=2)

add_body("Phase 1 (Months 1-6):", bold=True)
add_body("Direct sales to 25-50 companies in a single metro market. White-glove onboarding with data migration assistance. Goal: $5K-$15K MRR.")

add_body("Phase 2 (Months 7-12):", bold=True)
add_body("Content marketing (construction accounting guides, safety compliance checklists). Trade show presence (CONEXPO, World of Concrete). Referral program. Goal: $25K-$50K MRR.")

add_body("Phase 3 (Year 2+):", bold=True)
add_body("Channel partnerships with construction associations, accounting firms, and insurance brokers. Self-serve onboarding. Goal: $100K+ MRR.")

# ═══════════════════════════════════════════════════════════
# CONCLUSION
# ═══════════════════════════════════════════════════════════

add_heading_styled("7. Conclusion")

add_body("The construction industry's productivity problem is fundamentally a data integration problem. Companies waste 5-10% of project revenue reconciling data across disconnected systems, creating manual journal entries for automated business events, and chasing information that should be instantly available.")

add_body("Buildwrk solves this by eliminating the boundaries between project management, accounting, safety, equipment, property management, and business development. Every module reads from and writes to the same database, enforced by the same security model, and accessible through the same interface.")

add_body("With 80,000 lines of production TypeScript, 46 database tables, 461 security policies, 9 AI providers, and 19 feature modules, Buildwrk is a production-ready platform that delivers enterprise-grade capabilities at small-business prices.")

add_body("The question is not whether the construction industry will digitize \u2014 it's whether that digitization will happen through another generation of point solutions, or through a unified platform that treats construction data as an integrated whole.")

add_body("Buildwrk is building the latter.", bold=True, color=BLUE)

# Footer
doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("-" * 30)
run.font.color.rgb = MUTED

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Live Platform: construction-gamma-six.vercel.app")
run.font.size = Pt(10)
run.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Copyright 2026 Buildwrk. All rights reserved.")
run.font.size = Pt(9)
run.font.color.rgb = MUTED

# ═══════════════════════════════════════════════════════════
# SAVE DOCX
# ═══════════════════════════════════════════════════════════

output_dir = r"c:\Users\beltr\Construction.erp\docs\business"
docx_path = os.path.join(output_dir, "Buildwrk_White_Paper.docx")
doc.save(docx_path)
print(f"Saved DOCX: {docx_path}")


# ═══════════════════════════════════════════════════════════
# PDF GENERATION
# ═══════════════════════════════════════════════════════════

from fpdf import FPDF

class BuildwrkPDF(FPDF):
    def header(self):
        if self.page_no() > 1:
            self.set_font('Helvetica', 'B', 10)
            self.set_text_color(29, 78, 216)  # Blue
            self.cell(0, 8, 'Buildwrk', align='L')
            self.set_text_color(120, 113, 108)  # Muted
            self.set_font('Helvetica', '', 8)
            self.cell(0, 8, 'White Paper | February 2026', align='R')
            self.ln(4)
            self.set_draw_color(29, 78, 216)
            self.set_line_width(0.5)
            self.line(10, 14, 200, 14)
            self.ln(6)

    def footer(self):
        self.set_y(-15)
        self.set_font('Helvetica', '', 8)
        self.set_text_color(120, 113, 108)
        self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', align='C')

    def section_heading(self, text, level=1):
        self.ln(4)
        if level == 1:
            self.set_font('Helvetica', 'B', 18)
            self.set_text_color(29, 78, 216)
        elif level == 2:
            self.set_font('Helvetica', 'B', 14)
            self.set_text_color(29, 78, 216)
        else:
            self.set_font('Helvetica', 'B', 12)
            self.set_text_color(180, 83, 9)
        self.multi_cell(0, 8, text)
        self.ln(2)

    def body_text(self, text, bold=False):
        self.set_font('Helvetica', 'B' if bold else '', 10)
        self.set_text_color(41, 37, 36)
        self.multi_cell(0, 5.5, text)
        self.ln(2)

    def bullet(self, text):
        self.set_font('Helvetica', '', 10)
        self.set_text_color(41, 37, 36)
        x = self.get_x()
        self.cell(8, 5.5, "-")
        self.multi_cell(0, 5.5, text)
        self.ln(1)

    def add_table(self, headers, rows, col_widths=None):
        if col_widths is None:
            col_widths = [190 / len(headers)] * len(headers)

        # Header
        self.set_font('Helvetica', 'B', 8)
        self.set_fill_color(29, 78, 216)
        self.set_text_color(255, 255, 255)
        for i, header in enumerate(headers):
            self.cell(col_widths[i], 7, header, border=1, fill=True, align='C')
        self.ln()

        # Rows
        self.set_font('Helvetica', '', 8)
        self.set_text_color(41, 37, 36)
        for row in rows:
            for i, cell in enumerate(row):
                self.cell(col_widths[i], 6, cell, border=1)
            self.ln()
        self.ln(3)


pdf = BuildwrkPDF()
pdf.alias_nb_pages()
pdf.set_auto_page_break(auto=True, margin=20)

# ── Cover Page ──
pdf.add_page()
pdf.ln(50)
pdf.set_font('Helvetica', 'B', 42)
pdf.set_text_color(29, 78, 216)
pdf.cell(0, 15, 'Buildwrk', align='C')
pdf.ln(20)

# Amber bar
pdf.set_draw_color(180, 83, 9)
pdf.set_line_width(1)
pdf.line(80, pdf.get_y(), 130, pdf.get_y())
pdf.ln(10)

pdf.set_font('Helvetica', '', 18)
pdf.set_text_color(41, 37, 36)
pdf.cell(0, 10, 'The Unified Construction Intelligence Platform', align='C')
pdf.ln(15)

pdf.set_font('Helvetica', '', 14)
pdf.set_text_color(120, 113, 108)
pdf.cell(0, 8, 'A White Paper on Modernizing the', align='C')
pdf.ln(8)
pdf.cell(0, 8, '$1.3 Trillion Construction Industry', align='C')
pdf.ln(20)

pdf.set_font('Helvetica', 'B', 12)
pdf.set_text_color(180, 83, 9)
pdf.cell(0, 8, 'February 2026', align='C')
pdf.ln(30)

pdf.set_font('Helvetica', '', 9)
pdf.set_text_color(120, 113, 108)
pdf.cell(0, 6, 'CONFIDENTIAL', align='C')

# ── Content Pages ──
pdf.add_page()

pdf.section_heading("1. Executive Summary")
pdf.body_text("The construction industry generates $1.3 trillion in annual revenue in the United States alone, yet remains one of the least digitized sectors of the global economy. General contractors, developers, and property managers still rely on fragmented technology stacks --using one tool for project management, another for accounting, a third for safety compliance, and spreadsheets to bridge the gaps.")
pdf.body_text("Buildwrk is a unified construction ERP platform that consolidates project management, GAAP-compliant financial accounting, property management, safety compliance, equipment tracking, CRM, and AI-powered analytics into a single, multi-tenant SaaS application.")

pdf.section_heading("2. The Problem: Fragmented Technology")
pdf.section_heading("2.1 The Software Stack Tax", level=2)
pdf.body_text("A typical mid-market GC ($20M-$200M revenue) operates with 5-8 separate software systems:")

pdf.add_table(
    ["Function", "Common Tool", "Monthly Cost", "Data Silo"],
    [
        ["Project Management", "Procore", "$5K-$9K", "Schedules, RFIs"],
        ["Accounting", "Sage 300 CRE", "$3K-$8K", "GL, AP/AR, job costing"],
        ["Safety", "SafetyCulture", "$500-$2K", "Incidents, inspections"],
        ["HR/Payroll", "ADP/Paychex", "$1K-$3K", "Timesheets, wages"],
        ["CRM", "Salesforce", "$1.5K-$5K", "Leads, bids"],
    ],
    [55, 45, 40, 50]
)

pdf.body_text("Total: $12,500 --$31,500/month ($150K-$378K/year)", bold=True)

pdf.section_heading("2.2 The Mid-Market Vacuum", level=2)
pdf.body_text("Enterprise players serve firms above $500M with expensive implementations. Small-business tools serve firms below $5M with simplified features. The mid-market ($5M-$500M) is underserved.")

pdf.section_heading("3. The Solution: Unified Data Architecture")
pdf.section_heading("3.1 Single Source of Truth", level=2)
pdf.body_text("When a change order is approved, Buildwrk automatically:")
pdf.bullet("Updates the project budget (Project Management module)")
pdf.bullet("Creates a balanced journal entry (Financial module)")
pdf.bullet("Adjusts the contract value (Contract Management module)")
pdf.bullet("Logs the approval in the audit trail (Compliance module)")
pdf.bullet("Notifies the project manager and accountant")
pdf.body_text("All operations execute against the same database in a single transaction.")

pdf.section_heading("3.2 Multi-Tenant Security", level=2)
pdf.body_text("461 Row-Level Security (RLS) policies enforce complete data isolation between companies. 7 role levels enforce least-privilege access within each company.")

pdf.section_heading("3.3 GAAP-Compliant Financial Engine", level=2)
pdf.body_text("Retainage tracking: Separate JE lines for 5-10% holdback amounts.", bold=True)
pdf.body_text("Change order accounting: Auto-generated JEs based on CO type.", bold=True)
pdf.body_text("Job costing: CSI-code budget lines with variance analysis.", bold=True)
pdf.body_text("Financial statements: Income Statement, Balance Sheet, Cash Flow, Trial Balance --all auto-generated from the same journal entry data.")

pdf.section_heading("3.4 AI-Powered Intelligence", level=2)
pdf.body_text("9 LLM providers (OpenAI, Anthropic, Google, Groq, Mistral, Cohere, DeepSeek, xAI, AWS Bedrock) with function calling against live company data.")

pdf.add_page()
pdf.section_heading("4. Technical Architecture")

pdf.add_table(
    ["Layer", "Technology", "Rationale"],
    [
        ["Frontend", "Next.js 16.1.6", "Server-first rendering, SEO"],
        ["Language", "TypeScript (strict)", "Type safety, 646 files"],
        ["Database", "PostgreSQL (Supabase)", "ACID, RLS, real-time"],
        ["Auth", "Supabase Auth", "Email/password, OAuth"],
        ["Hosting", "Vercel", "Serverless, global CDN"],
        ["AI", "Vercel AI SDK", "Streaming, 9 providers"],
    ],
    [40, 55, 95]
)

pdf.body_text("Platform metrics: 80,000+ LOC | 46 tables | 461 RLS policies | 177 API endpoints | 178 pages | 35 migrations | 9 AI providers | 13 demo datasets", bold=True)

pdf.section_heading("5. Market Opportunity")
pdf.body_text("The construction management software market is valued at $11.58B in 2026, growing at 8.88% CAGR to $17.72B by 2031.")

pdf.add_table(
    ["Competitor", "Annual Cost (25 users)", "Focus", "Weakness"],
    [
        ["Procore", "$60K-$112K", "Project Mgmt", "No accounting"],
        ["Sage 300 CRE", "$36K-$96K", "Accounting", "Legacy UI"],
        ["Buildertrend", "$6K-$10K", "Residential", "Weak financials"],
        ["Buildwrk", "$3K-$6K", "Unified", "Building brand"],
    ],
    [40, 50, 45, 55]
)

pdf.section_heading("6. Business Model")
pdf.add_table(
    ["Tier", "Price", "Users", "Projects"],
    [
        ["Starter", "$99/mo", "5", "3 active"],
        ["Professional", "$249/mo", "25", "15 active"],
        ["Enterprise", "$499/mo", "Unlimited", "Unlimited"],
    ],
    [45, 50, 45, 50]
)

pdf.section_heading("7. Conclusion")
pdf.body_text("Buildwrk eliminates the boundaries between project management, accounting, safety, equipment, property management, and business development. Every module reads from and writes to the same database, enforced by the same security model.")
pdf.body_text("With 80,000 lines of production code, 46 database tables, 461 security policies, and 19 feature modules, Buildwrk delivers enterprise-grade capabilities at small-business prices.")

pdf.ln(10)
pdf.set_draw_color(180, 83, 9)
pdf.line(70, pdf.get_y(), 140, pdf.get_y())
pdf.ln(5)
pdf.set_font('Helvetica', 'B', 10)
pdf.set_text_color(29, 78, 216)
pdf.cell(0, 6, 'construction-gamma-six.vercel.app', align='C')
pdf.ln(6)
pdf.set_font('Helvetica', '', 8)
pdf.set_text_color(120, 113, 108)
pdf.cell(0, 5, 'Copyright 2026 Buildwrk. All rights reserved.', align='C')

# ── Save PDF ──
pdf_path = os.path.join(output_dir, "Buildwrk_White_Paper.pdf")
pdf.output(pdf_path)
print(f"Saved PDF: {pdf_path}")
